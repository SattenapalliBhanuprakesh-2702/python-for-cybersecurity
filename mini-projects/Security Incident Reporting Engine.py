
# prerequest : pip install pandas openpyxl python-docx reportlab

import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

try:
    df = pd.read_csv("auth_logs.csv")
except FileNotFoundError:
    print("CSV file not found")
    exit()
except Exception as e:
    print("Error reading CSV:", e)
    exit()

try:
    failed = df[df["status"] == "failed"]
    attackers = failed.groupby("ip").size().sort_values(ascending=False)
    attackers.to_csv("attack_summary.csv")
except Exception as e:
    print("Error processing data:", e)
    exit()

try:
    doc = Document()
    doc.add_heading("Incident Response Report", 1)
    doc.add_paragraph(f"Total Attempts: {len(df)}")
    doc.add_paragraph(f"Failed Attempts: {len(failed)}")
    for ip, count in attackers.items():
        doc.add_paragraph(f"{ip} → {count} failures")
    doc.save("incident_report.docx")
except Exception as e:
    print("Error creating Word report:", e)

try:
    c = canvas.Canvas("impact_report.pdf", pagesize=A4)
    c.drawString(50, 800, "Impact Report")
    y = 760
    for ip, count in attackers.items():
        c.drawString(50, y, f"{ip} → {count}")
        y -= 20
    c.save()
except Exception as e:
    print("Error creating PDF report:", e)

print("Incident reports generated")
