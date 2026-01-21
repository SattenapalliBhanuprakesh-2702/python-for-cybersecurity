from docx import Document
import pandas as pd 

df=pd.read_csv("auth_logs.csv")
failed=df[df["status"]=="failed"]

doc=Document()
doc.add_heading("Security Incident Report",level=1)
doc.add_paragraph(f"Total login attempts: {len(df)}")
doc.add_paragraph(f"Total failed attempts: {len(failed)}")

for ip,count in failed.groupby("ip").size().items():
    doc.add_paragraph(f"{ip} -> {count} failures")
doc.save("Security_report.docx")