import csv
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

def generate_csv_summary(logs, attacks):
    with open("reports/summary.csv", "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["IP", "Attempts"])
        for ip, c in attacks.items():
            w.writerow([ip, c])

def generate_word_report(logs, attacks):
    doc = Document()
    doc.add_heading("Incident Report", 1)
    doc.add_paragraph(f"Total logs: {len(logs)}")
    for ip, c in attacks.items():
        doc.add_paragraph(f"{ip} → {c} failures")
    doc.save("reports/incident_report.docx")

def generate_pdf_impact_report(logs, attacks, threat):
    c = canvas.Canvas("reports/impact_report.pdf", pagesize=A4)
    c.drawString(50, 800, "Impact Report")
    c.drawString(50, 780, f"Threat Score: {threat}")
    y = 760
    for ip, ccount in attacks.items():
        c.drawString(50, y, f"{ip} → {ccount}")
        y -= 20
    c.save()
